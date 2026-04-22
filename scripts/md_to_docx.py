"""Convert the carsales-clone-analysis.md file to a nicely formatted .docx.

Minimal Markdown handler tuned for this specific document:
- # / ## / ### headings
- bullet lists (- or *)
- numbered lists (1.)
- pipe tables (| a | b |)
- fenced code blocks (```)
- blockquotes (>)
- bold/italic inline (**x**, *x*, `x`)
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "carsales-clone-analysis.md"
DST = ROOT / "docs" / "carsales-clone-analysis.docx"


def add_inline(paragraph, text):
    """Parse **bold**, *italic*, `code` inline markup into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def style_body(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def main():
    doc = Document()
    style_body(doc)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 4))
            i += 1
            continue

        # Table: detect pipe table block
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-\s|:]+\|$", lines[i + 1].rstrip()):
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header + divider
            rows = []
            while i < len(lines) and lines[i].rstrip().startswith("|"):
                rows.append([c.strip() for c in lines[i].rstrip().strip("|").split("|")])
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
            tbl.style = "Light Grid Accent 1"
            for c, h in enumerate(header_cells):
                cell = tbl.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline(p, h)
                for run in p.runs:
                    run.bold = True
            for r_idx, row in enumerate(rows, start=1):
                for c_idx, val in enumerate(row):
                    if c_idx >= len(header_cells):
                        continue
                    cell = tbl.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    add_inline(cell.paragraphs[0], val)
            doc.add_paragraph()
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            i += 1
            continue

        # Bulleted list
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1))
            i += 1
            continue

        # Blank line
        if stripped == "":
            doc.add_paragraph()
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"wrote {DST}")


if __name__ == "__main__":
    main()
